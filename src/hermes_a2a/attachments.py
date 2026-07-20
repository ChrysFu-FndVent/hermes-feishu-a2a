from __future__ import annotations

import asyncio
import json
import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any, Protocol
from urllib.parse import unquote, urlsplit
from zipfile import ZipFile

from docx import Document
from pypdf import PdfReader

from .config import Settings
from .models import AttachmentReference, ExtractedAttachment

URL_CANDIDATE = re.compile(r"https?://[^\s\"'<>]+")
DRIVE_TOKEN_CHARS = frozenset(
    "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789_-"
)


class AttachmentError(RuntimeError):
    pass


@dataclass(frozen=True)
class DownloadedFile:
    name: str
    media_type: str
    content: bytes


class FeishuFileClient(Protocol):
    async def download_message_resource(
        self, message_id: str, file_key: str, *, max_bytes: int
    ) -> DownloadedFile: ...

    async def download_drive_file(self, file_token: str, *, max_bytes: int) -> DownloadedFile: ...


def parse_feishu_message(message: dict[str, Any]) -> tuple[str, list[AttachmentReference]]:
    raw_content = message.get("content", {})
    if isinstance(raw_content, str):
        try:
            parsed = json.loads(raw_content)
        except json.JSONDecodeError:
            parsed = {"text": raw_content}
    elif isinstance(raw_content, dict):
        parsed = raw_content
    else:
        parsed = {}

    text_parts: list[str] = []
    references: list[AttachmentReference] = []
    seen: set[tuple[str, str]] = set()

    if message.get("message_type") == "file":
        file_key = parsed.get("file_key")
        message_id = message.get("message_id")
        if isinstance(file_key, str) and isinstance(message_id, str):
            references.append(
                AttachmentReference(
                    kind="message_resource",
                    message_id=message_id,
                    file_key=file_key,
                    name=_clean_name(parsed.get("file_name")),
                )
            )
            seen.add(("message_resource", file_key))

    def walk(value: object) -> None:
        if isinstance(value, dict):
            href = value.get("href") or value.get("url")
            link_name = _clean_name(value.get("text"))
            if isinstance(href, str):
                _add_drive_references(href, link_name, references, seen)
            for key, item in value.items():
                if key in {"text", "title"} and isinstance(item, str):
                    cleaned = item.strip()
                    if cleaned:
                        text_parts.append(cleaned)
                    _add_drive_references(item, link_name, references, seen)
                elif key not in {"href", "url", "file_key", "file_name"}:
                    walk(item)
        elif isinstance(value, list):
            for item in value:
                walk(item)
        elif isinstance(value, str):
            _add_drive_references(value, None, references, seen)

    walk(parsed)
    return "\n".join(dict.fromkeys(text_parts)), references


def _add_drive_references(
    value: str,
    name: str | None,
    references: list[AttachmentReference],
    seen: set[tuple[str, str]],
) -> None:
    for match in URL_CANDIDATE.finditer(value):
        segments = urlsplit(match.group(0)).path.split("/")
        token: str | None = None
        for index, segment in enumerate(segments[:-1]):
            if segment != "file":
                continue
            token_chars: list[str] = []
            for character in segments[index + 1]:
                if character not in DRIVE_TOKEN_CHARS:
                    break
                token_chars.append(character)
            token = "".join(token_chars)
            break
        if not token:
            continue
        key = ("drive_file", token)
        if key in seen:
            continue
        references.append(AttachmentReference(kind="drive_file", file_token=token, name=name))
        seen.add(key)


def _clean_name(value: object) -> str | None:
    if not isinstance(value, str):
        return None
    cleaned = Path(unquote(value)).name.strip()
    return cleaned[:255] or None


class AttachmentIngestor:
    def __init__(self, settings: Settings, client: FeishuFileClient):
        self.settings = settings
        self.client = client

    async def resolve_all(self, references: list[AttachmentReference]) -> list[ExtractedAttachment]:
        attachments: list[ExtractedAttachment] = []
        total_chars = 0
        for reference in references:
            downloaded = await self._download(reference)
            preferred_name = reference.name or ""
            name = (
                preferred_name
                if Path(preferred_name).suffix.lower()
                in self.settings.feishu_file_allowed_extensions
                else downloaded.name
            )
            text = await asyncio.to_thread(
                extract_text,
                downloaded.content,
                name=name,
                media_type=downloaded.media_type,
                allowed_extensions=self.settings.feishu_file_allowed_extensions,
                max_chars=self.settings.feishu_file_max_extracted_chars,
                max_uncompressed_bytes=self.settings.feishu_file_max_uncompressed_bytes,
            )
            remaining = self.settings.feishu_file_max_total_chars - total_chars
            if remaining <= 0:
                raise AttachmentError("combined attachment text exceeds the configured limit")
            if len(text) > remaining:
                text = _truncate(text, remaining, "[truncated at combined attachment limit]")
            total_chars += len(text)
            attachments.append(
                ExtractedAttachment(
                    name=name,
                    media_type=downloaded.media_type,
                    text=text,
                    reference=reference,
                )
            )
        return attachments

    async def _download(self, reference: AttachmentReference) -> DownloadedFile:
        if reference.kind == "message_resource":
            if not reference.message_id or not reference.file_key:
                raise AttachmentError("message resource is missing its message ID or file key")
            return await self.client.download_message_resource(
                reference.message_id,
                reference.file_key,
                max_bytes=self.settings.feishu_file_max_bytes,
            )
        if not reference.file_token:
            raise AttachmentError("Drive file reference is missing its file token")
        return await self.client.download_drive_file(
            reference.file_token,
            max_bytes=self.settings.feishu_file_max_bytes,
        )


def extract_text(
    content: bytes,
    *,
    name: str,
    media_type: str,
    allowed_extensions: list[str],
    max_chars: int,
    max_uncompressed_bytes: int = 100 * 1024 * 1024,
) -> str:
    extension = Path(name).suffix.lower() or _extension_for_media_type(media_type)
    if extension not in allowed_extensions:
        raise AttachmentError(
            f"unsupported attachment type {extension or media_type}; "
            f"allowed: {', '.join(allowed_extensions)}"
        )
    try:
        if extension == ".pdf":
            reader = PdfReader(BytesIO(content))
            pages: list[str] = []
            extracted_chars = 0
            for page in reader.pages:
                page_text = page.extract_text() or ""
                pages.append(page_text)
                extracted_chars += len(page_text)
                if extracted_chars >= max_chars:
                    break
            text = "\n\n".join(pages)
        elif extension == ".docx":
            with ZipFile(BytesIO(content)) as archive:
                uncompressed_size = sum(item.file_size for item in archive.infolist())
            if uncompressed_size > max_uncompressed_bytes:
                raise AttachmentError(
                    "DOCX uncompressed content exceeds the configured safety limit"
                )
            document = Document(BytesIO(content))
            parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    parts.append("\t".join(cell.text.strip() for cell in row.cells))
            text = "\n".join(parts)
        else:
            text = content.decode("utf-8-sig", errors="replace")
    except AttachmentError:
        raise
    except Exception as exc:
        raise AttachmentError(f"failed to parse {name}: {exc}") from exc

    normalized = text.replace("\x00", "").strip()
    if not normalized:
        raise AttachmentError(f"no extractable text found in {name}")
    if len(normalized) > max_chars:
        return _truncate(normalized, max_chars, "[truncated at per-file extraction limit]")
    return normalized


def _extension_for_media_type(media_type: str) -> str:
    normalized = media_type.split(";", 1)[0].strip().lower()
    return {
        "application/pdf": ".pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "text/plain": ".txt",
        "text/markdown": ".md",
        "text/csv": ".csv",
        "application/json": ".json",
    }.get(normalized, "")


def _truncate(text: str, limit: int, note: str) -> str:
    if limit <= len(note) + 2:
        return text[:limit]
    return f"{text[: limit - len(note) - 2]}\n\n{note}"
