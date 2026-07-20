from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document

from hermes_a2a.attachments import AttachmentError, extract_text, parse_feishu_message


def test_parse_file_message_and_drive_links() -> None:
    text, references = parse_feishu_message(
        {
            "message_id": "om_file",
            "message_type": "post",
            "content": """{
              "title": "Reference files",
              "content": [[
                {"tag": "a", "text": "budget.docx", "href": "https://tenant.feishu.cn/file/token_one"},
                {"tag": "a", "text": "leasing.pdf", "href": "https://tenant.feishu.cn/file/token_two"},
                {"tag": "a", "text": "duplicate", "href": "https://tenant.feishu.cn/file/token_one"}
              ]]
            }""",
        }
    )

    assert "Reference files" in text
    assert [item.file_token for item in references] == ["token_one", "token_two"]
    assert [item.name for item in references] == ["budget.docx", "leasing.pdf"]


def test_parse_native_file_message() -> None:
    _, references = parse_feishu_message(
        {
            "message_id": "om_file",
            "message_type": "file",
            "content": '{"file_key":"file_key","file_name":"budget.docx"}',
        }
    )

    assert references[0].kind == "message_resource"
    assert references[0].message_id == "om_file"
    assert references[0].file_key == "file_key"


def test_parse_drive_link_without_regex_backtracking() -> None:
    hostile_prefix = "http://" * 10000
    _, references = parse_feishu_message(
        {"content": {"text": f"{hostile_prefix}/file/token_safe"}}
    )

    assert [item.file_token for item in references] == ["token_safe"]


def test_extract_docx_paragraphs_and_tables() -> None:
    document = Document()
    document.add_paragraph("Budget assumptions")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Rent"
    table.cell(0, 1).text = "100000"
    buffer = BytesIO()
    document.save(buffer)

    text = extract_text(
        buffer.getvalue(),
        name="budget.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        allowed_extensions=[".docx"],
        max_chars=10000,
    )

    assert "Budget assumptions" in text
    assert "Rent\t100000" in text


def test_extract_pdf_text() -> None:
    text = extract_text(
        _minimal_pdf("Lease assumptions"),
        name="leasing.pdf",
        media_type="application/pdf",
        allowed_extensions=[".pdf"],
        max_chars=10000,
    )

    assert "Lease assumptions" in text


def test_reject_unsupported_file_type() -> None:
    with pytest.raises(AttachmentError, match="unsupported attachment type"):
        extract_text(
            b"legacy",
            name="legacy.doc",
            media_type="application/msword",
            allowed_extensions=[".pdf", ".docx"],
            max_chars=10000,
        )


def _minimal_pdf(text: str) -> bytes:
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        f"<< /Length {len(stream)} >>\nstream\n".encode() + stream + b"\nendstream",
    ]
    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{index} 0 obj\n".encode() + obj + b"\nendobj\n")
    xref = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode())
    output.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode()
    )
    return bytes(output)
